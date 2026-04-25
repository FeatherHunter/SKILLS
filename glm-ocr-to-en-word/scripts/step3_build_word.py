#!/usr/bin/env python3
"""阶段3：读取layout_details JSON，用python-docx构建Word文档（逐页模式）

用法：
  python step3_build_word.py <input.json> <output.docx>
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import json
import os
import sys
import re
import logging


class GLMOCRToWord:

    DEFAULT_CONFIG = {
        "font_cn": "宋体",
        "font_en": "Calibri",
        "title_font_size": 16,
        "heading_font_size": 14,
        "body_font_size": 12,
        "table_font_size": 10,
        "note_font_size": 9,
        "page_margins_cm": 2.54,
        "image_max_width_inches": 6.0,
        "line_spacing": 1.15,
    }

    def __init__(self, config=None):
        self.config = {**self.DEFAULT_CONFIG, **(config or {})}

    def convert(self, layout_details, output_path, logger=None):
        """主转换函数"""
        doc = self._create_document()
        total_elements = 0
        table_count = 0
        image_count = 0

        for page_idx, page_elements in enumerate(layout_details):
            if page_idx > 0:
                doc.add_page_break()

            for element in page_elements:
                label = element.get("label", "text")
                if label == "table":
                    table_count += 1
                elif label == "image":
                    image_count += 1
                self._add_element(doc, element)
                total_elements += 1

        doc.save(output_path)
        if logger:
            logger.info(f"统计: {len(layout_details)} 页, {total_elements} 元素, {table_count} 表格, {image_count} 图片")
        else:
            print(f"Word文档已保存: {output_path}")
            print(f"处理统计: {len(layout_details)} 页, {total_elements} 个元素")

    def _create_document(self):
        """创建Word文档并设置全局样式"""
        doc = Document()

        margin = Cm(self.config["page_margins_cm"])
        for section in doc.sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin

        style = doc.styles['Normal']
        font = style.font
        font.name = self.config["font_en"]
        font.size = Pt(self.config["body_font_size"])
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.config["font_cn"])

        style.paragraph_format.line_spacing = self.config["line_spacing"]

        return doc

    def _add_element(self, doc, element):
        """根据元素类型添加到Word文档"""
        label = element.get("label", "text")
        content = element.get("content", "")
        bbox = element.get("bbox_2d", [])

        if label == "text":
            self._add_text(doc, content, bbox)
        elif label == "table":
            self._add_table(doc, content)
        elif label == "image":
            self._add_image(doc, content)
        elif label == "formula":
            self._add_formula(doc, content)
        else:
            self._add_text(doc, content, bbox)

    def _parse_html_content(self, content):
        """
        解析HTML内容，提取纯文本和对齐方式
        关键修复：处理<div align="center">等HTML标签，而非直接插入
        """
        if not content:
            return "", WD_ALIGN_PARAGRAPH.LEFT

        content = content.strip()
        
        # 检查是否是HTML（简单检测）
        has_html_tags = bool(re.search(r'<\w+[^>]*>', content))
        
        if has_html_tags:
            soup = BeautifulSoup(content, "html.parser")
            text = soup.get_text(strip=True)
            
            # 检测对齐方式
            alignment = WD_ALIGN_PARAGRAPH.LEFT
            div = soup.find('div')
            if div and div.get('align') == 'center':
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif div and div.get('align') == 'right':
                alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 检查是否有标题标签
            h_tag = soup.find(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
            if h_tag:
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 处理Markdown标题标记
            if text.startswith('## '):
                text = text[3:]
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif text.startswith('# '):
                text = text[2:]
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            return text, alignment
        
        # 非HTML，处理Markdown标题
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        if content.startswith('## '):
            content = content[3:]
            alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif content.startswith('# '):
            content = content[2:]
            alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return content, alignment

    def _add_text(self, doc, content, bbox=None):
        """
        添加文本段落
        关键修复：使用_parse_html_content解析HTML，而非直接插入原始字符串
        """
        if not content or not content.strip():
            return

        # 解析HTML获取纯文本和对齐方式（修复HTML标签显示问题）
        text, alignment = self._parse_html_content(content)
        
        if not text:
            return

        font_size, is_bold, _ = self._infer_text_style(text, bbox)
        is_heading = font_size >= self.config["heading_font_size"]

        if is_heading:
            level = 1 if font_size >= self.config["title_font_size"] else 2
            p = doc.add_heading(text, level=level)
            p.alignment = alignment
            for run in p.runs:
                run.font.name = self.config["font_en"]
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config["font_cn"])
        else:
            p = doc.add_paragraph()
            p.alignment = alignment
            run = p.add_run(text)
            run.font.name = self.config["font_en"]
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config["font_cn"])
            run.font.size = Pt(font_size)
            run.bold = is_bold

    def _infer_text_style(self, text, bbox=None):
        """基于bbox坐标推断样式"""
        font_size = self.config["body_font_size"]
        is_bold = False
        alignment = WD_ALIGN_PARAGRAPH.LEFT

        if bbox and len(bbox) == 4:
            x1, y1, x2, y2 = bbox

            if y1 < 0.2 and (x2 - x1) > 0.5:
                is_bold = True
                if y1 < 0.1 and len(text) < 30:
                    font_size = self.config["title_font_size"]
                    alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    font_size = self.config["heading_font_size"]
            elif y2 > 0.9:
                font_size = self.config["note_font_size"]

        elif len(text) < 30 and not text.endswith(
            ('.', ',', ';', ':', '。', '，', '；', '：', ')', '）')
        ):
            is_bold = True
            if len(text) < 15:
                font_size = self.config["title_font_size"]
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                font_size = self.config["heading_font_size"]

        return font_size, is_bold, alignment

    def _add_table(self, doc, html_content):
        """
        将HTML表格转换为Word表格
        使用BeautifulSoup解析HTML结构，然后用python-docx构建
        """
        if not html_content or not html_content.strip():
            return

        try:
            soup = BeautifulSoup(html_content, "html.parser")
            tables = soup.find_all("table")

            for table_html in tables:
                self._process_single_table(doc, table_html)
                
        except Exception as e:
            # 解析失败时，将原始内容作为文本插入
            p = doc.add_paragraph(f"[Table parse error: {str(e)[:50]}]")
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            doc.add_paragraph(html_content[:200])  # 插入前200字符供调试

    def _process_single_table(self, doc, table_html):
        """处理单个HTML表格"""
        rows_html = table_html.find_all("tr")
        if not rows_html:
            return

        # 构建二维网格
        grid = []
        for row_idx, row_html in enumerate(rows_html):
            cells_html = row_html.find_all(["td", "th"])

            while len(grid) <= row_idx:
                grid.append([])

            col_idx = 0
            for cell_html in cells_html:
                colspan = int(cell_html.get("colspan", 1))
                rowspan = int(cell_html.get("rowspan", 1))

                while col_idx < len(grid[row_idx]) and grid[row_idx][col_idx] is not None:
                    col_idx += 1

                needed_cols = col_idx + colspan
                for r in range(len(grid)):
                    while len(grid[r]) < needed_cols:
                        grid[r].append(None)

                for r in range(row_idx, row_idx + rowspan):
                    while len(grid) <= r:
                        grid.append([])
                    while len(grid[r]) < needed_cols:
                        grid[r].append(None)
                    for c in range(col_idx, col_idx + colspan):
                        if r == row_idx and c == col_idx:
                            grid[r][c] = {
                                "cell_html": cell_html,
                                "origin": (row_idx, col_idx),
                                "colspan": colspan,
                                "rowspan": rowspan,
                            }
                        else:
                            grid[r][c] = {
                                "cell_html": None,
                                "origin": (row_idx, col_idx),
                                "colspan": 0,
                                "rowspan": 0,
                            }

                col_idx += colspan

        num_rows = len(grid)
        num_cols = max(len(row) for row in grid) if grid else 0

        if num_rows == 0 or num_cols == 0:
            return

        for row in grid:
            while len(row) < num_cols:
                row.append(None)

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # 填充文本
        for r in range(num_rows):
            for c in range(num_cols):
                cell_info = grid[r][c]
                if cell_info and cell_info["cell_html"] is not None:
                    cell_text = cell_info["cell_html"].get_text(strip=True)
                    target_cell = table.cell(r, c)
                    target_cell.text = cell_text

                    for paragraph in target_cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = self.config["font_en"]
                            run._element.rPr.rFonts.set(
                                qn('w:eastAsia'), self.config["font_cn"]
                            )
                            run.font.size = Pt(self.config["table_font_size"])

        # 合并单元格
        merged_origins = set()

        for r in range(num_rows):
            for c in range(num_cols):
                cell_info = grid[r][c]
                if not cell_info or cell_info["cell_html"] is None:
                    continue

                origin = cell_info["origin"]
                if origin in merged_origins:
                    continue
                merged_origins.add(origin)

                cs = cell_info["colspan"]
                rs = cell_info["rowspan"]

                if cs > 1 or rs > 1:
                    try:
                        start_cell = table.cell(origin[0], origin[1])
                        end_cell = table.cell(origin[0] + rs - 1, origin[1] + cs - 1)
                        start_cell.merge(end_cell)

                        merged_text = cell_info["cell_html"].get_text(strip=True)
                        start_cell.text = merged_text
                        for paragraph in start_cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.name = self.config["font_en"]
                                run._element.rPr.rFonts.set(
                                    qn('w:eastAsia'), self.config["font_cn"]
                                )
                                run.font.size = Pt(self.config["table_font_size"])
                    except Exception as e:
                        pass

        doc.add_paragraph()

    def _add_image(self, doc, image_source):
        """添加图片到Word文档"""
        if not image_source:
            return

        try:
            if os.path.exists(image_source):
                doc.add_picture(
                    image_source,
                    width=Inches(self.config["image_max_width_inches"])
                )
            elif image_source.startswith("http"):
                import requests
                from io import BytesIO
                response = requests.get(image_source, timeout=30)
                doc.add_picture(
                    BytesIO(response.content),
                    width=Inches(self.config["image_max_width_inches"])
                )
            else:
                p = doc.add_paragraph(f"[Image: {image_source}]")
                p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        except Exception as e:
            p = doc.add_paragraph(f"[Image loading failed: {str(e)[:50]}]")
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

    def _add_formula(self, doc, content):
        """添加公式（保留为文本）"""
        if not content:
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(content)
        run.font.name = "Cambria Math"
        run.font.size = Pt(self.config["body_font_size"])
        run.italic = True


def setup_logging(temp_dir):
    """配置日志"""
    logs_dir = os.path.join(temp_dir, "logs")
    os.makedirs(logs_dir, exist_ok=True)
    log_file = os.path.join(logs_dir, "step3_build_word.log")
    
    for handler in logging.getLogger().handlers[:]:
        logging.getLogger().removeHandler(handler)
    
    logging.basicConfig(
        level=logging.INFO,
        format='[%(asctime)s] [%(levelname)s] %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8', mode='a'),
            logging.StreamHandler(sys.stdout)
        ]
    )
    return logging.getLogger(__name__)


def main():
    if len(sys.argv) < 3:
        print("用法: python step3_build_word.py <input.json> <output.docx>")
        sys.exit(1)

    json_path = sys.argv[1]
    output_path = sys.argv[2]

    if not os.path.exists(json_path):
        print(f"错误: 文件不存在: {json_path}")
        sys.exit(1)

    output_dir = os.path.dirname(os.path.abspath(output_path))
    temp_dir = os.path.normpath(os.path.join(output_dir, ".."))
    logger = setup_logging(temp_dir)
    
    basename = os.path.basename(json_path)
    logger.info(f"读取: {basename}")
    
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    if isinstance(data, dict) and "elements" in data:
        page_num = data.get("page_number", "?")
        elements = data["elements"]
        layout_details = [elements]
        logger.info(f"页码: {page_num}，元素数: {len(elements)}")
    elif isinstance(data, list):
        layout_details = data
        logger.info(f"总页数: {len(layout_details)}")
    else:
        layout_details = [data]
    
    converter = GLMOCRToWord()
    converter.convert(layout_details, output_path, logger)
    
    file_size = os.path.getsize(output_path) / 1024
    logger.info(f"保存: {output_path} ({file_size:.1f} KB)")


if __name__ == "__main__":
    main()
